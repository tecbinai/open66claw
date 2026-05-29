"""
生成《医药大健康·短视频引流全自动化系统·业务需求方案》Word 文档
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 颜色常量 ──────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1A, 0x56, 0xDB)
BLUE_LIGHT = RGBColor(0x3B, 0x82, 0xF6)
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)
GRAY_MED = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xDC, 0x26, 0x26)
GREEN = RGBColor(0x16, 0xA3, 0x4A)

# ── 辅助函数 ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "D1D5DB")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None):
    """添加美化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, '1A56DB')

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = GRAY_DARK
            bg = 'F8FAFC' if r_idx % 2 == 0 else 'FFFFFF'
            set_cell_shading(cell, bg)

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph('')  # 间距
    return table

def add_heading_styled(doc, text, level=1):
    """添加带装饰的标题"""
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run('█ ')
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(18)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = BLUE_DARK
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # 下划线装饰
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(8)
        r = line.add_run('━' * 60)
        r.font.color.rgb = RGBColor(0xDB, 0xE4, 0xF0)
        r.font.size = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run('▎')
        run.font.color.rgb = BLUE_LIGHT
        run.font.size = Pt(14)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = GRAY_DARK
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('● ')
        run.font.color.rgb = BLUE_LIGHT
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_DARK
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_body(doc, text):
    """正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY_DARK
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_note(doc, text):
    """提示/注释"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = GRAY_MED
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_flow(doc, text):
    """流程图文字（等宽居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE_DARK
    run.font.name = 'Consolas'
    return p

def add_bullet(doc, text, indent=0):
    """项目符号"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.8)
    marker = '•' if indent == 0 else '◦'
    run = p.add_run(f'{marker} {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_DARK
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('医药大健康')
run.font.bold = True
run.font.size = Pt(36)
run.font.color.rgb = BLUE_DARK
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('短视频引流全自动化系统')
run.font.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE_LIGHT
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0xDB, 0xE4, 0xF0)
run.font.size = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run('业 务 需 求 方 案')
run.font.size = Pt(20)
run.font.color.rgb = GRAY_MED
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(4):
    doc.add_paragraph('')

# 封面信息表
info_items = [
    ('客户行业', '医药 / 大健康'),
    ('引流平台', '抖音 · 快手 · 小红书'),
    ('业务模式', '短视频/图文引流 → 私域沉淀 → 成交转化'),
    ('文档版本', 'v1.0'),
    ('日期', '2026 年 3 月 19 日'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_MED
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_DARK
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 分页
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '目  录', 1)

toc_items = [
    '一、业务全景',
    '二、M1 爆款素材采集',
    '三、M2 素人账号采集',
    '四、M3 内容自动化生产',
    '五、M4 评论区截留引流',
    '六、M5 点赞引流',
    '七、M6 关注引流',
    '八、M7 私信引流转化',
    '九、M8 账号矩阵管理',
    '十、M9 私域承接与成交',
    '十一、全链路转化漏斗',
    '十二、医药行业合规红线',
    '十三、系统功能模块总览',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE_DARK
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 一、业务全景
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '一、业务全景', 1)

add_heading_styled(doc, '1.1 核心业务链路', 2)
add_flow(doc, '找爆款  →  建知识库  →  做内容  →  矩阵分发  →  截留引流  →  私域成交')

add_heading_styled(doc, '1.2 系统架构总览', 2)
add_styled_table(doc,
    ['模块编号', '模块名称', '核心职责'],
    [
        ['M1', '爆款素材采集', '找到行业内正在爆的视频/图文，提供选题和素材'],
        ['M2', '素人账号采集', '找到真实素人账号，下载内容用于二创和养号参考'],
        ['M3', '内容自动化生产', '基于知识库+模板库，批量生产差异化内容'],
        ['M4', '评论区截留引流', '在热门视频评论区抢占前排，引导回访主页'],
        ['M5', '点赞引流', '批量点赞目标用户作品，触发回访'],
        ['M6', '关注引流', '三步关注法，利用回关机制建立粉丝关系'],
        ['M7', '私信引流转化', '对高意向用户一对一沟通，引导加微信'],
        ['M8', '账号矩阵管理', '多账号分工协作，养号+风控+状态监控'],
        ['M9', '私域承接与成交', '微信/企微沉淀，持续运营至成交'],
    ],
    col_widths=[2.5, 4, 10]
)

add_heading_styled(doc, '1.3 数据流向', 2)
add_flow(doc, '爆款采集 + 素人采集  →  素材库  →  内容生产  →  待发布内容库  →  矩阵分发')
add_flow(doc, '↓')
add_flow(doc, '三大平台发布  →  评论/点赞/关注/私信引流  →  用户加微信  →  私域成交')
add_flow(doc, '↓')
add_flow(doc, '数据反馈  →  优化各环节')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 二、M1 爆款素材采集
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '二、M1 爆款素材采集', 1)

add_heading_styled(doc, '2.1 业务目的', 2)
add_body(doc, '解决"发什么能火"的问题。快速找到行业内正在爆的视频和图文，提取选题方向、内容结构、热门标签，作为内容生产的输入源。')

add_heading_styled(doc, '2.2 业务流程', 2)
add_flow(doc, '维护关键词库  →  三平台并发搜索  →  按爆款条件筛选  →  列表预览')
add_flow(doc, '→  批量下载（视频/图文/文案/封面）→  入库分类打标签  →  爆款拆解分析')

add_heading_styled(doc, '2.3 功能需求清单', 2)
add_styled_table(doc,
    ['编号', '功能', '说明'],
    [
        ['1.1', '关键词库管理', '维护行业关键词（养生/调理/中医/免疫力/睡眠/食疗…），支持增删改查'],
        ['1.2', '三平台统一搜索', '输入关键词，同时搜索抖音、快手、小红书，结果统一展示'],
        ['1.3', '爆款条件筛选', '支持按点赞量、评论量、收藏量、发布时间、博主粉丝范围等条件过滤'],
        ['1.4', '结果列表预览', '缩略图 + 标题 + 核心数据（赞/评/藏），支持批量勾选'],
        ['1.5', '一键批量下载', '去水印下载视频/图文/文案/封面/BGM 信息'],
        ['1.6', '素材库管理', '按平台、话题、类型、热度自动分类打标签，支持搜索和筛选'],
        ['1.7', '爆款拆解', '记录爆款的开头钩子、内容结构、话术风格、标签组合，形成可复用的选题卡'],
    ],
    col_widths=[1.5, 3.5, 11.5]
)

add_heading_styled(doc, '2.4 爆款筛选条件参考', 2)
add_styled_table(doc,
    ['筛选条件', '抖音', '快手', '小红书'],
    [
        ['点赞量', '≥ 3,000', '≥ 2,000', '≥ 1,000'],
        ['评论量', '≥ 500', '≥ 300', '≥ 200'],
        ['收藏量', '—', '—', '≥ 500（核心指标）'],
        ['发布时间', '最近 72 小时', '最近 72 小时', '最近 7 天'],
        ['点赞率', '≥ 3%', '≥ 5%', '—'],
        ['博主粉丝范围', '可自定义', '可自定义', '可自定义'],
        ['关键词命中', '标题/文案/标签', '标题/文案/标签', '标题/文案/标签'],
    ],
    col_widths=[3.5, 4, 4, 5]
)
add_note(doc, '所有阈值均支持用户自定义调整。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 三、M2 素人账号采集
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '三、M2 素人账号采集', 1)

add_heading_styled(doc, '3.1 业务目的', 2)
add_body(doc, '找到真实的普通人账号，下载其主页信息及历史作品。用途有二：一是作为二次创作的素材来源（真实感强），二是学习素人的表达风格用于养号。')

add_heading_styled(doc, '3.2 业务流程', 2)
add_flow(doc, '设定筛选条件  →  平台搜索  →  系统自动判定是否素人  →  一键下载  →  归档入库')

add_heading_styled(doc, '3.3 素人账号判定规则', 2)
add_styled_table(doc,
    ['维度', '符合条件', '排除条件'],
    [
        ['粉丝量', '100 ~ 5,000', '< 100（僵尸号）或 > 5,000'],
        ['认证状态', '无蓝V / 无黄V', '有任何认证标识'],
        ['商业痕迹', '无橱窗 / 无商品链接 / 无星图', '有任何商业化功能'],
        ['内容风格', '生活化、非专业制作', '明显包装感或专业剪辑'],
        ['发布频率', '不规律', '固定日更（像营销号）'],
        ['账号年龄', '≥ 3 个月', '近期新注册'],
    ],
    col_widths=[3, 6.5, 7]
)

add_heading_styled(doc, '3.4 采集内容', 2)
add_styled_table(doc,
    ['采集项', '说明'],
    [
        ['主页信息', '头像、昵称、简介、粉丝数、获赞数、IP 属地'],
        ['全部作品', '视频文件 + 封面 + 文案 + 标签 + 各项数据（赞/评/藏/转）'],
        ['评论样本（可选）', '该账号在其他视频下发过的评论，用于学习真实评论风格'],
    ],
    col_widths=[4, 12.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 四、M3 内容自动化生产
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '四、M3 内容自动化生产', 1)

add_heading_styled(doc, '4.1 业务目的', 2)
add_body(doc, '解决"怎么批量做出能火的内容"的问题。基于爆款素材、专业知识库、素人素材，系统化生产短视频和图文内容，并生成多个差异化版本供矩阵账号使用。')

add_heading_styled(doc, '4.2 三大输入源', 2)
add_styled_table(doc,
    ['输入源', '提供内容', '核心价值'],
    [
        ['① 爆款素材库', '热门选题、爆款结构模板、热门 BGM/标签', '告诉你"什么形式能火"'],
        ['② 专业知识库', '医药知识、养生食疗、中医科普、合规话术', '告诉你"说什么内容可信"'],
        ['③ 素人素材库', '素人作品、真实表达风格、生活化场景', '告诉你"怎么说像真人"'],
    ],
    col_widths=[3.5, 7, 6]
)

add_heading_styled(doc, '4.3 专业知识库（核心壁垒）', 2)
add_body(doc, '这是医药行业区别于其他行业的核心竞争力。知识库的深度和专业度直接决定了内容的可信度和转化效果。')

add_heading_styled(doc, '知识库内容结构', 3)
add_styled_table(doc,
    ['分类', '内容', '来源', '示例'],
    [
        ['养生科普', '日常调理、食疗方、节气养生', '中医教材/公开文献', '春季养肝三个习惯'],
        ['成分百科', '草本植物、食材功效、营养成分', '药典/营养学教材', '黄芪的 5 个日常用法'],
        ['生活方式', '作息、运动、饮食、情绪管理', '健康指南/WHO 建议', '失眠人群的睡前 30 分钟'],
        ['常见误区', '大众认知错误、伪科学辟谣', '权威医学机构', '这 3 个养胃方法其实是错的'],
        ['时令热点', '换季、流感季、节气对应内容', '时事+专业结合', '倒春寒，这类人要特别注意'],
        ['合规话术', '安全用词替换表、禁用词清单', '广告法/平台规则', '"调理"替代"治疗"'],
    ],
    col_widths=[2.5, 5, 4, 5]
)

add_heading_styled(doc, '知识库管理要求', 3)
add_bullet(doc, '标签化：每条知识打标签（症状/人群/季节/场景），方便检索匹配')
add_bullet(doc, '合规审核：入库前审核，确保不含违规表述')
add_bullet(doc, '定期更新：跟随季节、热点、平台规则变化更新')
add_bullet(doc, '权威溯源：每条知识标注来源出处，可追溯')

add_heading_styled(doc, '4.4 爆款模板库', 2)

add_heading_styled(doc, '视频模板', 3)
add_styled_table(doc,
    ['模板类型', '结构', '示例标题'],
    [
        ['恐惧开头型', '抛出问题→放大痛点→给出方案→引导关注', '还在这样喝水？你的肾在求救'],
        ['反常识型', '颠覆认知→解释原因→正确做法', '红枣补血？大错特错'],
        ['清单型', '数字标题→逐条讲解→引导收藏', '中医推荐的 5 个泡脚方'],
        ['故事型', '个人经历→遇到问题→如何解决→现状', '我妈失眠 10 年，后来…'],
        ['对比型', '错误做法 vs 正确做法', '90% 的人枸杞都吃错了'],
        ['热点借势型', '热点事件→关联健康话题→专业解读', 'XX 事件背后的健康真相'],
    ],
    col_widths=[3, 6.5, 7]
)

add_heading_styled(doc, '图文模板（小红书为主）', 3)
add_styled_table(doc,
    ['模板类型', '结构', '示例'],
    [
        ['教程卡片', '封面标题图 + 3~6 张步骤图 + 文案', '祛湿茶配方｜每天一杯'],
        ['知识图谱', '一张信息量大的科普长图', '9 种体质自测表'],
        ['合集种草', '多方法/多好物合集', '打工人养胃好物合集'],
        ['问答型', '模拟用户提问→专业解答', '经常熬夜怎么补救？'],
    ],
    col_widths=[3, 7, 6.5]
)

add_heading_styled(doc, '4.5 内容生产完整流程', 2)
add_styled_table(doc,
    ['步骤', '名称', '输入', '输出'],
    [
        ['S1', '选题生成', '爆款库+知识库+时令热点', '选题卡'],
        ['S2', '脚本生成', '选题卡+爆款模板+知识库', '脚本初稿'],
        ['S3', '合规审核', '脚本+敏感词库', '审核通过的脚本'],
        ['S4', '素材匹配', '脚本内容', '匹配好的视频片段/图片/BGM'],
        ['S5', '成片生成', '脚本+素材+配音', '视频成片 / 图文卡片'],
        ['S6', '差异化处理', '原始成片', '多个差异化版本'],
        ['S7', '入库排期', '多版本成片', '按账号/平台/时段分配'],
        ['S8', '数据回收', '发布后48h数据', '反馈优化选题和模板'],
    ],
    col_widths=[1.5, 3, 5.5, 6.5]
)

add_heading_styled(doc, '选题生成权重', 3)
add_styled_table(doc,
    ['选题来源', '权重', '逻辑'],
    [
        ['爆款库近 3 天热门选题', '40%', '已验证有流量，成功率高'],
        ['专业知识库匹配', '30%', '有专业深度，建立信任'],
        ['时令/节气/热点', '20%', '借势流量，时效性强'],
        ['竞品账号分析', '10%', '参考同行表现，查漏补缺'],
    ],
    col_widths=[5, 2, 9.5]
)

add_heading_styled(doc, '合规审核 — 敏感词替换表', 3)
add_styled_table(doc,
    ['禁用词', '安全替换'],
    [
        ['治疗 / 治愈', '调理 / 改善'],
        ['药 / 保健品', '好物 / 好方法'],
        ['疗效显著', '感觉不错 / 舒服多了'],
        ['患者 / 病人', '朋友 / 家人'],
        ['根治 / 痊愈', '好转 / 有变化'],
        ['降压 / 降糖', '注意饮食习惯'],
        ['处方 / 医嘱', '建议 / 经验分享'],
    ],
    col_widths=[5, 11.5]
)

add_heading_styled(doc, '差异化处理维度', 3)
add_body(doc, '同一选题/脚本生成多个版本，供矩阵内不同账号使用，避免平台判定搬运：')
add_styled_table(doc,
    ['差异化维度', '做法'],
    [
        ['封面', '换文案/配色/排版'],
        ['开头', '换不同的钩子句'],
        ['BGM', '换不同的背景音乐'],
        ['配音', '换音色/语速'],
        ['画面', '镜像翻转/调色/换素材片段'],
        ['文案', '同义改写，保持意思不变'],
        ['标签', '换不同的标签组合'],
    ],
    col_widths=[4, 12.5]
)

add_heading_styled(doc, '4.6 内容分发排期', 2)
add_styled_table(doc,
    ['维度', '规则'],
    [
        ['发布时段', '早 7~9 点 / 午 12~14 点 / 晚 19~22 点（流量高峰）'],
        ['发布频率', '内容号：每天 1~2 条；引流号：每周 2~3 条'],
        ['平台适配', '同一选题发不同平台时，调整时长/格式/标签'],
        ['矩阵错开', '同一内容的不同版本，错开至少 2 小时发布'],
        ['热点响应', '突发热点 2 小时内出内容，抢时效'],
    ],
    col_widths=[3, 13.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 五、M4 评论区截留引流
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '五、M4 评论区截留引流', 1)

add_heading_styled(doc, '5.1 业务目的', 2)
add_body(doc, '在他人热门视频评论区抢占靠前位置，通过有价值的评论吸引用户点击头像→回访主页→进入私域。这是整个引流体系的核心动作。')

add_heading_styled(doc, '5.2 业务流程', 2)
add_flow(doc, '【准备】目标视频筛选 + 话术库维护 + 账号池管理')
add_flow(doc, '↓')
add_flow(doc, '【执行】生成任务（视频+话术+账号匹配）→ 观看视频 → 发布评论 → 小号互顶')
add_flow(doc, '↓')
add_flow(doc, '【监控】评论状态监控 → 引流效果追踪 → 数据反馈优化')

add_heading_styled(doc, '5.3 目标视频筛选规则', 2)
add_styled_table(doc,
    ['维度', '标准'],
    [
        ['行业相关性', '视频内容与医药/健康/养生直接相关'],
        ['评论区活跃度', '评论量 ≥ 200，且有真实讨论'],
        ['博主粉丝量', '10 万 ~ 100 万（中腰部最佳）'],
        ['发布时间', '优先 30 分钟内的新视频（抢前排黄金窗口）'],
        ['数据趋势', '数据正在上升中的视频'],
    ],
    col_widths=[4, 12.5]
)

add_heading_styled(doc, '5.4 截留时机', 2)
add_styled_table(doc,
    ['时机', '视频发布后', '效果', '说明'],
    [
        ['黄金窗口', '0 ~ 30 分钟', '★★★★★', '评论少，早期评论容易推到前排'],
        ['优质窗口', '30 分钟 ~ 2 小时', '★★★★', '配合小号互顶可以上热评'],
        ['常规窗口', '2 ~ 6 小时', '★★★', '需要靠高质量"神评论"突围'],
        ['长尾窗口', '6 小时以上', '★★', '仅针对持续爆火的视频'],
    ],
    col_widths=[3, 4, 3, 6.5]
)

add_heading_styled(doc, '5.5 话术分层体系', 2)
add_styled_table(doc,
    ['层级', '名称', '占比', '目的', '示例'],
    [
        ['L1', '纯互动', '60%', '养号权重', '"太对了，深有同感"'],
        ['L2', '价值补充', '20%', '展示专业度', '"补充一下，XX 情况还可以…"'],
        ['L3', '软引导', '15%', '引发好奇→回访', '"之前也有这困扰，后来找到了方法"'],
        ['L4', '钩子型', '5%', '直接收集意向', '"整理了份 XX 指南，需要的扣 1"'],
    ],
    col_widths=[1.5, 2.5, 1.5, 4, 7]
)

add_heading_styled(doc, '话术管理要求', 3)
add_bullet(doc, '每层级储备 ≥ 30 条话术，支持增删改查')
add_bullet(doc, '支持变量替换（随机 emoji、语气词、口语化表达）增加自然感')
add_bullet(doc, '同一视频下不出现重复话术')
add_bullet(doc, '同一账号连续评论不用相似句式')

add_heading_styled(doc, '5.6 执行节奏', 2)
add_styled_table(doc,
    ['参数', '抖音', '快手', '小红书'],
    [
        ['单号日评论上限', '15 ~ 20 条', '20 ~ 30 条', '10 ~ 15 条'],
        ['评论间隔', '1 ~ 3 分钟随机', '1 ~ 3 分钟随机', '2 ~ 4 分钟随机'],
        ['活跃时段', '8:00 ~ 23:00', '8:00 ~ 23:00', '8:00 ~ 23:00'],
        ['休息机制', '每 3~5 条休息 5~15 分钟', '同左', '同左'],
        ['异常熔断', '检测到验证码立即暂停', '同左', '同左'],
    ],
    col_widths=[4, 4, 4, 4.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 六、M5 点赞引流
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '六、M5 点赞引流', 1)

add_heading_styled(doc, '6.1 业务目的', 2)
add_body(doc, '通过批量点赞目标用户的作品，触发平台通知→用户好奇回访主页→被主页内容吸引→关注或私信。')

add_heading_styled(doc, '6.2 业务流程', 2)
add_flow(doc, '筛选目标用户  →  对每人连续点赞 3~5 条作品  →  用户回访主页  →  引导关注/私信')

add_heading_styled(doc, '6.3 运营参数', 2)
add_styled_table(doc,
    ['项目', '参数'],
    [
        ['单号日操作量', '200 ~ 500 次'],
        ['预期回访率', '1 ~ 3%'],
        ['风险等级', '低（点赞是平台鼓励的正常行为）'],
        ['核心前提', '主页必须优化好——简介有钩子、置顶有价值内容'],
    ],
    col_widths=[4, 12.5]
)

# ═══════════════════════════════════════════════════════════
# 七、M6 关注引流
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '七、M6 关注引流', 1)

add_heading_styled(doc, '7.1 业务目的', 2)
add_body(doc, '批量关注目标用户，利用社交礼仪性回关心理，建立粉丝关系后再通过私信或内容做进一步转化。')

add_heading_styled(doc, '7.2 业务流程', 2)
add_flow(doc, '筛选目标  →  三步关注法（点赞→评论→关注）→  等待回关  →  回关者进入私信转化')

add_heading_styled(doc, '7.3 运营参数', 2)
add_styled_table(doc,
    ['项目', '参数'],
    [
        ['单号日操作量', '50 ~ 100 次'],
        ['预期回关率', '三步关注法：10 ~ 15%（直接关注仅 3 ~ 5%）'],
        ['风险等级', '中（平台有日关注上限，超过会限流）'],
        ['适用场景', '冷启动阶段快速建立粉丝基础'],
    ],
    col_widths=[4, 12.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 八、M7 私信引流转化
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '八、M7 私信引流转化', 1)

add_heading_styled(doc, '8.1 业务目的', 2)
add_body(doc, '对已有互动基础的高意向用户进行一对一私信沟通，引导添加微信进入私域。这是从公域到私域的关键转化环节。')

add_heading_styled(doc, '8.2 触发场景', 2)
add_styled_table(doc,
    ['优先级', '触发场景', '说明'],
    [
        ['最高', '用户主动私信', '已有明确意向，必须立即响应'],
        ['高', '用户回复了你的评论', '有互动基础，趁热跟进'],
        ['中', '用户回关了你', '有一定兴趣，发送欢迎语'],
        ['低', '用户点赞了你的内容', '轻度兴趣，可尝试触达'],
    ],
    col_widths=[2.5, 5, 9]
)

add_heading_styled(doc, '8.3 四步转化话术', 2)
add_styled_table(doc,
    ['步骤', '目的', '话术示例'],
    [
        ['第 1 步', '打招呼 + 确认需求', '"看到你在 XX 视频下的评论，是想了解 XX 方面的信息吗？"'],
        ['第 2 步', '提供价值 + 建立信任', '"我之前也研究过这个问题，有一些经验可以分享…"'],
        ['第 3 步', '引导转移 + 给出理由', '"平台私信经常吞消息，方便的话加个微信沟通更顺畅"'],
        ['第 4 步', '发送联系方式', '图片形式/谐音/引导看主页简介（避免直接发微信号）'],
    ],
    col_widths=[2, 4.5, 10]
)

add_heading_styled(doc, '8.4 运营参数', 2)
add_styled_table(doc,
    ['项目', '新号', '老号'],
    [
        ['单号日私信量', '5 ~ 10 条', '20 ~ 30 条'],
        ['最佳转化窗口', '互动后 24 小时内', '—'],
        ['预期回复率', '10 ~ 20%', '15 ~ 25%'],
        ['回复→加微信', '20 ~ 40%', '20 ~ 40%'],
        ['风险等级', '高', '中高'],
    ],
    col_widths=[5, 5.5, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 九、M8 账号矩阵管理
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '九、M8 账号矩阵管理', 1)

add_heading_styled(doc, '9.1 业务目的', 2)
add_body(doc, '多账号分工协作，分散单号风险，放大整体引流效果。建议矩阵总规模 20~50 个账号。')

add_heading_styled(doc, '9.2 矩阵架构', 2)
add_styled_table(doc,
    ['角色', '数量', '职责'],
    [
        ['品牌主号', '1 个', '承接流量，展示品牌形象，不直接做引流动作'],
        ['评论截留号', '5 ~ 10 个', '在别人热门视频下评论引流'],
        ['内容号', '3 ~ 5 个', '发布知识科普/案例故事'],
        ['互动号', '5 ~ 10 个', '给截留评论点赞上热评 + 给内容号制造热度'],
        ['点赞关注号', '5 ~ 10 个', '批量点赞、关注拉新'],
    ],
    col_widths=[3, 2.5, 11]
)

add_heading_styled(doc, '9.3 养号生命周期', 2)
add_styled_table(doc,
    ['阶段', '天数', '操作', '禁止事项'],
    [
        ['注册完善', 'D1 ~ 3', '完善资料、实名认证、绑定手机', '不发内容、不评论'],
        ['冷启动', 'D4 ~ 7', '每天刷 30~60 分钟，点赞 10~20 个，评论 2~3 条', '不带任何营销信息'],
        ['活跃期', 'D8 ~ 14', '发 1~2 条生活化原创，持续互动', '不放联系方式'],
        ['测试期', 'D15 ~ 21', '发 3~5 条垂直内容，观察数据', '数据异常需排查'],
        ['正式运营', 'D22+', '按分工执行引流任务', '遵守各项操作上限'],
    ],
    col_widths=[2.5, 2.5, 6.5, 5]
)

add_heading_styled(doc, '9.4 封号红线', 2)
add_styled_table(doc,
    ['高危行为', '后果'],
    [
        ['同一设备登录多个账号', '关联封号'],
        ['评论/私信中出现微信号、电话、链接', '删评/禁言/封号'],
        ['短时间大量重复操作', '限流/降权'],
        ['新号直接发广告内容', '永久封号'],
        ['搬运他人内容不做改动', '限流/下架'],
        ['使用模拟器或多开软件', '设备级封禁'],
    ],
    col_widths=[8, 8.5]
)

add_heading_styled(doc, '9.5 降权信号', 2)
add_bullet(doc, '视频播放量持续低于 200')
add_bullet(doc, '评论/私信发出后对方看不到')
add_bullet(doc, '搜索自己的昵称找不到账号')
add_bullet(doc, '发布内容长时间显示"审核中"')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 十、M9 私域承接与成交
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '十、M9 私域承接与成交', 1)

add_heading_styled(doc, '10.1 业务目的', 2)
add_body(doc, '将从公域引流过来的用户沉淀到微信私域，通过持续运营建立信任，最终完成成交转化。')

add_heading_styled(doc, '10.2 承接方式', 2)
add_styled_table(doc,
    ['形式', '适用场景'],
    [
        ['个人微信号', '高客单价产品，需要一对一深度信任'],
        ['企业微信', '规模化运营，合规需求高'],
        ['微信社群', '批量运营，需要氛围带动成交'],
        ['公众号/小程序', '内容沉淀 + 标准化商品成交'],
    ],
    col_widths=[4, 12.5]
)

add_heading_styled(doc, '10.3 私域运营节奏', 2)
add_styled_table(doc,
    ['时间', '动作', '目的'],
    [
        ['Day 0', '欢迎语 + 自我介绍 + 赠送"行业资料"', '建立好感'],
        ['Day 1~3', '朋友圈展示专业内容（科普、案例）', '展示专业度'],
        ['Day 3~7', '私聊跟进，了解具体需求', '建立信任'],
        ['Day 7~14', '推荐产品/服务，给出方案', '促成转化'],
        ['Day 14+', '持续朋友圈运营，定期回访', '长期经营'],
    ],
    col_widths=[2.5, 7, 7]
)

add_heading_styled(doc, '10.4 核心原则', 2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
run = p.add_run('⚠ 公域平台（抖音/快手/小红书）只做知识输出和信任建立，所有产品推荐和成交动作只在私域完成。')
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RED
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 十一、全链路转化漏斗
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '十一、全链路转化漏斗', 1)

add_heading_styled(doc, '11.1 单次转化预估', 2)
add_styled_table(doc,
    ['环节', '数据', '转化率'],
    [
        ['评论/点赞/关注曝光触达', '10,000 次', '—'],
        ['主页访问', '100 ~ 300 次', '1 ~ 3%'],
        ['产生互动（关注/私信/回复）', '15 ~ 45 次', '5 ~ 15%'],
        ['添加微信', '3 ~ 18 人', '20 ~ 40%'],
        ['深度咨询', '1 ~ 9 人', '30 ~ 50%'],
        ['成交', '0.3 ~ 2.7 单', '10 ~ 30%'],
    ],
    col_widths=[6, 4, 6.5]
)

add_heading_styled(doc, '11.2 矩阵产能预估', 2)
add_styled_table(doc,
    ['指标', '数值'],
    [
        ['单号日产能', '评论 15 条 + 点赞 300 次 + 关注 50 次'],
        ['20 号矩阵日触达', '约 7,000 ~ 10,000 次'],
        ['月预估进粉', '90 ~ 540 人'],
        ['月预估成交', '9 ~ 80 单（取决于客单价和私域转化能力）'],
    ],
    col_widths=[5, 11.5]
)
add_note(doc, '以上为保守估算，实际效果取决于内容质量、话术水平、账号权重、私域运营能力等因素。')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 十二、医药行业合规红线
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '十二、医药行业合规红线', 1)

add_heading_styled(doc, '12.1 内容合规', 2)
add_styled_table(doc,
    ['✅ 可以做', '❌ 绝对不能做'],
    [
        ['健康科普、养生常识、生活方式建议', '疾病治疗承诺、功效保证'],
        ['个人调理经历分享（第一人称叙述）', '处方药任何形式的推广'],
        ['食疗/运动/作息建议', '代替就医的暗示'],
        ['原创科普内容', '伪造用户案例、虚假证书'],
        ['引用公开学术研究（注明来源）', '前后效果对比图'],
        ['"调理/改善/缓解/建议"', '"治愈/根治/100%有效/国家级"'],
    ],
    col_widths=[8.25, 8.25]
)

add_heading_styled(doc, '12.2 引流合规', 2)
add_styled_table(doc,
    ['✅ 可以做', '❌ 绝对不能做'],
    [
        ['主页简介间接引导', '评论/私信中直接放微信号、电话、链接'],
        ['评论区输出知识价值、引发互动', '评论区硬广、群发广告'],
        ['通过内容质量吸引自然关注', '使用机器人刷量'],
    ],
    col_widths=[8.25, 8.25]
)

add_heading_styled(doc, '12.3 法律底线', 2)
add_styled_table(doc,
    ['法规', '核心要求'],
    [
        ['《广告法》第 16~18 条', '医疗/药品/保健品广告不得含有功效断言、代言推荐'],
        ['《药品管理法》', '处方药严禁面向公众广告；药品广告须经省级药监审批'],
        ['《互联网广告管理办法》', '互联网广告须标注"广告"；素人软广同样受监管'],
        ['《个人信息保护法》', '不得违规采集用户个人信息'],
    ],
    col_widths=[5, 11.5]
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(8)
run = p.add_run('⚠ 违法后果：罚款 20~100 万元，严重可追究刑事责任。')
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RED
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_heading_styled(doc, '12.4 合规策略总结', 2)
add_bullet(doc, '身份定位：做"健康生活方式分享者"，不做"卖药的"')
add_bullet(doc, '内容策略：80% 纯知识科普 + 20% 软性方向引导，不提具体产品')
add_bullet(doc, '转化分离：公域只做信任建立，私域才做产品推荐和成交')
add_bullet(doc, '素材二创：只借鉴选题和结构，重新组织内容，不搬运不抄袭')
add_bullet(doc, '合规前置：所有内容发布前过敏感词库审核，宁可少说不可说错')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 十三、系统功能模块总览
# ═══════════════════════════════════════════════════════════
add_heading_styled(doc, '十三、系统功能模块总览', 1)

add_styled_table(doc,
    ['模块', '核心功能', '使用角色'],
    [
        ['M1 爆款采集', '关键词搜索 + 条件筛选 + 批量下载 + 素材库', '内容运营'],
        ['M2 素人采集', '素人识别 + 主页下载 + 风格分析', '内容运营'],
        ['M3 内容生产', '知识库 + 模板库 + 选题→脚本→成片→差异化', '内容制作'],
        ['M4 评论截留', '目标筛选 + 话术匹配 + 自动评论 + 效果监控', '引流运营'],
        ['M5 点赞引流', '目标筛选 + 批量点赞 + 数据追踪', '引流运营'],
        ['M6 关注引流', '三步关注 + 回关监控 + 取关清理', '引流运营'],
        ['M7 私信转化', '触发规则 + 话术模板 + 发送节奏控制', '转化运营'],
        ['M8 账号矩阵', '养号排期 + 状态监控 + 风控预警 + 封号记录', '账号管理'],
        ['M9 私域承接', '微信/企微沉淀 + 社群运营 + 成交管理', '销售/客服'],
        ['数据看板', '全链路转化数据 + 内容效果排行 + ROI 分析', '管理层'],
    ],
    col_widths=[3.5, 8, 5]
)

# ── 尾页 ──────────────────────────────────────────────────
doc.add_page_break()
for _ in range(8):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('━' * 40)
r.font.color.rgb = RGBColor(0xDB, 0xE4, 0xF0)
r.font.size = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run('本文档为业务需求梳理，不涉及具体技术实现方案。')
run.font.size = Pt(11)
run.font.color.rgb = GRAY_MED
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('各模块的技术选型、开发排期、成本预算将在需求确认后另行输出。')
run.font.size = Pt(11)
run.font.color.rgb = GRAY_MED
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 保存 ──────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), '医药大健康-短视频引流全自动化-业务需求方案.docx')
doc.save(output_path)
print('Word doc generated: ' + output_path)
